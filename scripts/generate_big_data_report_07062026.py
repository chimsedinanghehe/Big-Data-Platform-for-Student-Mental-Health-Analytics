from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "Bao_cao_big_data.07062026.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])


def add_toc(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = 'TOC \\o "1-3" \\h \\z \\u'
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "Cập nhật mục lục trong Microsoft Word bằng Update Field."
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text, end])


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    header = table.rows[0]
    set_repeat_table_header(header)
    for index, value in enumerate(headers):
        cell = header.cells[index]
        cell.text = value
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, "DCEAF2")
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = str(value)
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    if widths:
        for row in table.rows:
            for index, width in enumerate(widths):
                row.cells[index].width = Inches(width)
    doc.add_paragraph()
    return table


def add_bullets(doc: Document, items: list[str], level: int = 0) -> None:
    style = "List Bullet" if level == 0 else "List Bullet 2"
    for item in items:
        doc.add_paragraph(item, style=style)


def add_numbers(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_code(doc: Document, code: str) -> None:
    for line in code.strip("\n").splitlines():
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.right_indent = Inches(0.15)
        p.paragraph_format.space_after = Pt(0)
        set_cell = OxmlElement("w:shd")
        set_cell.set(qn("w:fill"), "F2F4F5")
        p._p.get_or_add_pPr().append(set_cell)
        run = p.add_run(line or " ")
        run.font.name = "Consolas"
        run.font.size = Pt(8.5)
    doc.add_paragraph()


def add_note(doc: Document, title: str, text: str, color: str = "EAF4F8") -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, color)
    p = cell.paragraphs[0]
    run = p.add_run(f"{title}: ")
    run.bold = True
    p.add_run(text)
    doc.add_paragraph()


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def build_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.68)
    add_page_number(section.footer.paragraphs[0])

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    for name, size, color in [
        ("Title", 24, "173A4D"),
        ("Heading 1", 17, "173A4D"),
        ("Heading 2", 14, "216B82"),
        ("Heading 3", 11.5, "2C5968"),
    ]:
        styles[name].font.name = "Arial"
        styles[name].font.size = Pt(size)
        styles[name].font.color.rgb = RGBColor.from_string(color)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(80)
    run = title.add_run("BÁO CÁO TRIỂN KHAI VÀ VẬN HÀNH\nNỀN TẢNG BIG DATA PHÂN TÍCH SỨC KHỎE TÂM LÝ")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(23, 58, 77)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("Survey, Chat Logs, GCS Data Lake, Kafka, Dataproc Serverless Spark và Dashboard").bold = True
    doc.add_paragraph()
    add_table(
        doc,
        ["Thông tin", "Giá trị"],
        [
            ["Tên file", OUTPUT.name],
            ["Ngày lập báo cáo", "07/06/2026 (Asia/Saigon)"],
            ["Project GCP", "student-mental-health-496205"],
            ["Region xử lý", "asia-southeast1"],
            ["GCS Data Lake", "gs://student-mental-health-lake-nhom1-2026"],
            ["Phạm vi", "Các công việc triển khai, tối ưu và kiểm thử ngày 06-07/06/2026"],
        ],
        [1.7, 4.7],
    )
    add_note(
        doc,
        "Nguyên tắc phạm vi",
        "Backend, Frontend và Streamlit Dashboard có thể chạy tại máy phát triển; toàn bộ Spark xử lý dữ liệu được chạy trên Dataproc Serverless và đọc/ghi GCS. Không chạy Spark local.",
    )
    doc.add_page_break()

    add_heading(doc, "Mục lục", 1)
    add_toc(doc.add_paragraph())
    doc.add_page_break()

    add_heading(doc, "1. Tóm tắt điều hành", 1)
    doc.add_paragraph(
        "Hệ thống đã được tổ chức theo kiến trúc Data Lake phân tầng Bronze → Silver → Gold trên Google Cloud Storage. "
        "Survey từ ứng dụng được hợp nhất thành một snapshot Parquet duy nhất, unique theo user_id. Chat logs được gửi "
        "qua Kafka trên Compute Engine, gom theo batch rồi ghi JSONL vào Bronze. Dataproc Serverless chạy Spark để chuẩn hóa, "
        "ẩn danh, khử trùng lặp, xây dựng các chỉ số phân tích và cập nhật Dashboard."
    )
    add_bullets(
        doc,
        [
            "Survey: không tạo một file nhỏ cho mỗi lần submit; sử dụng bronze/app_survey_snapshot/survey_all.parquet.",
            "Chat: Backend producer → Kafka topic student-chat-logs → consumer duy nhất → GCS Bronze JSONL.",
            "Spark: sử dụng Dataproc Serverless, AQE bật, shuffle/default parallelism = 8, output partitions = 4.",
            "Gold survey dùng versioned run_id; Dashboard tự chọn run đầy đủ và mới nhất.",
            "Dashboard Tổng quan gộp Học sinh và Sinh viên; các tab chi tiết lọc theo nhóm.",
            "Production mode giảm Spark action; count/schema/sample/null report chỉ bật khi debug hoặc audit.",
        ],
    )
    add_table(
        doc,
        ["Kết quả kiểm thử cuối", "Giá trị"],
        [
            ["Batch tài khoản mới nhất", "100 tài khoản: 50 Học sinh + 50 Sinh viên"],
            ["Survey API thành công", "100/100, 0 lỗi"],
            ["Chat Kafka thành công", "100/100, 0 lỗi"],
            ["Thời gian gửi Survey + Chat", "327,155 giây với 5 worker"],
            ["Bronze survey snapshot", "307 dòng, 307 user_id duy nhất, 0 duplicate"],
            ["Gold survey hiện tại", "286.280 tổng; 20.256 Học sinh; 266.024 Sinh viên"],
            ["Tỷ lệ nguy cơ hiện tại", "44,129174%; Học sinh 46,5739%; Sinh viên 43,9430%"],
            ["Chat Gold ngày 07/06/2026", "100 tin nhắn; Học sinh 50; Sinh viên 50"],
            ["Dataproc batches", "4/4 SUCCEEDED"],
            ["Health check cuối", "Backend, Frontend, Dashboard, Kafka tunnel, Bronze và Gold đều healthy"],
        ],
        [2.4, 4.0],
    )

    add_heading(doc, "2. Công cụ và nền tảng sử dụng", 1)
    add_table(
        doc,
        ["Thành phần", "Công cụ", "Vai trò"],
        [
            ["Frontend", "Vite + React", "Đăng ký, đăng nhập, khảo sát dạng wizard/radio-card, chatbot, liên kết Dashboard."],
            ["Backend", "FastAPI + Python", "Xác thực, profile, survey API, chat/RAG API, producer Kafka, snapshot worker."],
            ["CSDL nghiệp vụ", "PostgreSQL trong Ubuntu/WSL", "Lưu user, profile, trạng thái survey và survey response."],
            ["Event streaming", "Apache Kafka 3.9.1 KRaft trên Compute Engine VM", "Nhận chat event từ Backend và cung cấp cho consumer GCS."],
            ["Data Lake", "Google Cloud Storage", "Lưu Bronze, Silver, Gold, scripts và dữ liệu lịch sử."],
            ["Xử lý Big Data", "Apache Spark trên Dataproc Serverless", "Chuẩn hóa, dedup, tính feature, tổng hợp Gold."],
            ["Dashboard", "Streamlit + Pandas + Plotly", "Đọc Gold mới nhất và hiển thị Tổng quan/Học sinh/Sinh viên."],
            ["Quản trị", "PowerShell, Bash, Google Cloud APIs/gcloud, SSH", "Khởi chạy dịch vụ, tunnel, submit batch, kiểm tra log."],
            ["Định dạng", "Parquet + Snappy; JSONL cho event Bronze", "Schema ổn định, đọc Spark nhanh, giảm kích thước và small files."],
        ],
        [1.25, 2.0, 3.2],
    )

    add_heading(doc, "3. Kiến trúc tổng thể", 1)
    add_code(
        doc,
        """
Người dùng Web
  ├─ Submit Survey
  │    └─ FastAPI → PostgreSQL → single snapshot worker
  │         └─ GCS Bronze survey_all.parquet
  │              └─ Dataproc Survey Bronze→Silver→Gold
  │                   └─ Streamlit Dashboard
  │
  └─ Gửi Chat
       └─ FastAPI producer → SSH tunnel → Kafka VM
            └─ Kafka consumer duy nhất → GCS Bronze JSONL
                 └─ Dataproc Chat Bronze→Silver→Gold
                      └─ Streamlit Dashboard
""",
    )
    add_note(
        doc,
        "Không trộn sai mẫu số",
        "Tỷ lệ survey tính theo người tham gia. Chat metrics tính theo tin nhắn/phiên. Tổng quan hiển thị cả hai nguồn nhưng không cộng chat message vào mẫu số survey.",
        "FFF3CD",
    )

    add_heading(doc, "4. Trình tự công việc đã thực hiện", 1)
    add_numbers(
        doc,
        [
            "Kiểm tra hai bộ historical standardized: school_survey_standardized.csv và university_survey_standardized.csv; giữ chúng làm nền dữ liệu lịch sử.",
            "Xây dựng survey API và luồng một user chỉ hoàn thành survey một lần.",
            "Xây dựng giao diện khảo sát riêng: câu hỏi có dấu, radio-card, bắt buộc trả lời mới sang câu tiếp theo, lưu draft localStorage, khóa Dashboard trước khi hoàn thành.",
            "Giữ canonical answer ở Backend/Spark nhưng thay nhãn hiển thị Sinh viên thành câu trả lời có ý nghĩa; bộ Học sinh giữ cấu trúc.",
            "Xây dựng snapshot survey Parquet duy nhất, deduplicate theo user_id và validate file tạm trước khi cập nhật file chính.",
            "Tối ưu Survey Bronze→Silver và Silver→Gold: production/debug tách biệt, giảm action và partition, bật AQE, versioned Gold.",
            "Cấu hình Kafka VM và consumer chat ghi GCS Bronze theo batch.",
            "Bổ sung metadata tuổi, giới tính, learner_type, lớp/khóa, audience_group và survey status vào chat event.",
            "Tối ưu Chat Bronze→Silver và Silver→Gold để Dashboard Tổng quan/Học sinh/Sinh viên đọc được.",
            "Nâng cấp Kafka producer với idempotence, acks=all, retry, delivery timeout và key ổn định theo phiên.",
            "Bổ sung script tự kiểm tra/khôi phục Kafka tunnel và health check toàn bộ luồng production.",
            "Bổ sung Survey Gold current manifest để Dashboard đọc nhanh run Gold hiện hành, tránh quét toàn bộ prefix.",
            "Chạy 100 tài khoản survey + chat qua luồng production; 100/100 survey và 100/100 chat thành công.",
            "Submit bốn Dataproc batch và xác nhận đều SUCCEEDED.",
            "Đối chiếu Gold trước/sau và chỉnh độ chính xác KPI Tổng quan từ 2 lên 3 chữ số thập phân.",
        ],
    )

    add_heading(doc, "5. Luồng Survey", 1)
    add_heading(doc, "5.1 Quy tắc nghiệp vụ", 2)
    add_bullets(
        doc,
        [
            "Mỗi user chỉ được hoàn thành survey đúng một lần.",
            "survey_type được xác định theo profile: school hoặc university.",
            "Frontend giữ draft theo user_id và survey_type; draft bị xóa khi submit thành công.",
            "Backend validate đáp án theo option canonical; FE chỉ thay đổi nhãn hiển thị.",
            "Snapshot Bronze phải unique theo user_id và chứa tối thiểu user_id, age, survey_type, submitted_at, schema_version.",
        ],
    )
    add_heading(doc, "5.2 Ghi snapshot an toàn", 2)
    add_numbers(
        doc,
        [
            "Single worker lấy các survey_response chưa export từ PostgreSQL.",
            "Dùng PostgreSQL advisory lock survey_snapshot_writer để không có hai worker ghi đồng thời.",
            "Đọc survey_all.parquet hiện tại từ GCS.",
            "Merge bản ghi mới và drop_duplicates theo user_id.",
            "Validate schema bắt buộc, user_id rỗng, duplicate user_id và survey_type.",
            "Ghi bytes Parquet vào object tạm survey_all_tmp_<uuid>.parquet.",
            "Đọc lại object tạm và validate lần nữa.",
            "Chỉ sau khi validate thành công mới upload/replace survey_all.parquet.",
            "Nếu có lỗi, xóa file tạm và giữ nguyên snapshot chính.",
            "Đánh dấu response đã export trong PostgreSQL sau khi snapshot được cập nhật.",
        ],
    )
    add_note(
        doc,
        "Mức chịu lỗi survey",
        "Single writer + advisory lock ngăn race condition ở worker; temporary object + validate bảo vệ file chính; Silver tiếp tục dedup để tạo lớp phòng vệ thứ hai.",
    )

    add_heading(doc, "6. Luồng Chat và Kafka", 1)
    add_heading(doc, "6.1 Kafka VM hiện tại", 2)
    add_table(
        doc,
        ["Cấu hình", "Giá trị"],
        [
            ["VM", "student-chat-streaming-m; region/zone asia-southeast1-a"],
            ["Kafka", "3.9.1; Scala 2.13; KRaft single-node"],
            ["Topic", "student-chat-logs"],
            ["Partitions", "4"],
            ["Replication factor", "1"],
            ["Retention", "168 giờ"],
            ["Segment size", "268.435.456 bytes"],
            ["Kafka heap", "-Xms512m -Xmx1g"],
            ["Consumer group", "gcs-log-creators-v2"],
            ["Consumer batch", "Tối đa 50 event hoặc chờ tối đa 15 giây"],
            ["GCS output", "bronze/chat_logs/date=YYYY-MM-DD/kafka_batch_*.jsonl"],
        ],
        [2.0, 4.4],
    )
    add_heading(doc, "6.2 Metadata chat", 2)
    add_bullets(
        doc,
        [
            "event_id, event_type, timestamp, anonymous_session_id và user_id_hash.",
            "user_age, user_gender, learner_type, grade, class_level, user_group/audience_group.",
            "survey_type và survey_completed.",
            "question, answer và standalone_query đã mask PII.",
            "model, is_document_rag, emotion và safety.",
        ],
    )
    add_heading(doc, "6.3 Commit và chịu lỗi", 2)
    add_bullets(
        doc,
        [
            "Consumer tắt auto commit.",
            "Event được gom vào buffer; ghi JSONL lên GCS trước.",
            "Chỉ commit offset đồng bộ sau khi upload GCS thành công.",
            "systemd Restart=always cho consumer và Restart=on-failure cho Kafka.",
            "Backend producer flush với timeout và báo lỗi rõ nếu broker/tunnel không truy cập được.",
        ],
    )
    add_note(
        doc,
        "Giới hạn hiện tại",
        "Kafka đang single-node và replication factor = 1, vì vậy chịu lỗi ở mức tiến trình/consumer nhưng chưa HA ở mức broker. Production quan trọng nên dùng Kafka cluster nhiều broker hoặc managed Kafka.",
        "FFF3CD",
    )

    add_heading(doc, "7. Kiến trúc phân tầng dữ liệu GCS", 1)
    add_code(
        doc,
        """
gs://student-mental-health-lake-nhom1-2026/
├─ bronze/
│  ├─ survey_standardized/
│  │  ├─ school_survey_standardized.csv
│  │  └─ university_survey_standardized.csv
│  ├─ app_survey_snapshot/
│  │  └─ survey_all.parquet
│  └─ chat_logs/
│     └─ date=YYYY-MM-DD/
│        └─ kafka_batch_<timestamp>_<uuid>.jsonl
├─ silver/
│  ├─ survey_cleaned/
│  ├─ survey_cleaned_invalid/
│  ├─ anonymized_chat/
│  └─ anonymized_chat_invalid/
├─ gold/
│  └─ dashboard_tables/
│     ├─ survey_analytic_features/run_id=<run_id>/
│     ├─ survey_overview_summary/run_id=<run_id>/
│     ├─ survey_response_by_date/run_id=<run_id>/
│     ├─ survey_demographic_summary/run_id=<run_id>/
│     ├─ survey_question_distribution/run_id=<run_id>/
│     ├─ survey_numeric_summary/run_id=<run_id>/
│     ├─ chat_hourly_metrics/
│     ├─ chat_risk_summary/
│     ├─ chat_topic_summary/
│     ├─ chat_construct_summary/
│     └─ chat_model_usage/
├─ scripts/
│  ├─ survey_pipeline/
│  ├─ chat_pipeline/
│  ├─ kafka/
│  ├─ scheduler/
│  └─ e2e/
└─ tmp/
   └─ survey_gold_work/
""",
    )
    add_table(
        doc,
        ["Tầng", "Mục đích", "Nguyên tắc"],
        [
            ["Bronze", "Dữ liệu nguồn gần nguyên bản", "Không tính dashboard; giữ khả năng replay/backfill."],
            ["Silver", "Chuẩn hóa, cast, dedup, ẩn danh, metadata", "Không đọc ngược Bronze trong Silver→Gold."],
            ["Gold", "Bảng feature và summary phục vụ Dashboard", "Chỉ tạo từ Silver; summary nhỏ, đọc nhanh."],
        ],
        [1.1, 2.5, 2.8],
    )

    add_heading(doc, "8. Cấu hình Spark/Dataproc Serverless", 1)
    add_table(
        doc,
        ["Thông số production", "Giá trị", "Lý do"],
        [
            ["dynamicAllocation", "false", "Dữ liệu khoảng 286k dòng; cấu hình cố định giảm thời gian điều phối."],
            ["executor instances", "2", "Đủ song song nhưng không cấp tài nguyên quá mức."],
            ["executor cores", "4", "Tổng 8 executor cores phù hợp parallelism 8."],
            ["executor memory", "8g", "Đủ xử lý wide survey và aggregate."],
            ["driver cores/memory", "4 / 8g", "Đủ lập kế hoạch và metadata."],
            ["spark.default.parallelism", "8", "Không tạo quá nhiều task cho dữ liệu nhỏ/medium."],
            ["spark.sql.shuffle.partitions", "8", "Giảm shuffle task và small files so với 24/48."],
            ["AQE", "enabled", "Cho Spark coalesce partition sau shuffle theo dữ liệu thực tế."],
            ["AQE coalesce partitions", "enabled", "Tự giảm partition thừa."],
            ["output partitions", "4", "Cân bằng song song và số file trên GCS."],
            ["Gold summary files", "coalesce(1)", "Summary nhỏ; Dashboard đọc nhanh."],
            ["Gold analytic features", "coalesce(4)", "Tránh một file quá lớn và tránh small files."],
            ["FileOutputCommitter", "algorithm.version=2", "Giảm overhead commit/rename trên object storage."],
            ["Temp Gold stage", "disabled mặc định", "Tránh ghi/đọc GCS trung gian cho dữ liệu hiện tại."],
        ],
        [2.0, 1.4, 3.0],
    )

    add_heading(doc, "9. Chiến lược tối ưu Spark", 1)
    add_heading(doc, "9.1 Giảm Spark actions", 2)
    add_bullets(
        doc,
        [
            "Production không chạy count/show/printSchema/null count/groupBy chỉ để log.",
            "Không đọc lại Silver/Gold từ GCS sau write trừ khi bật --enable-output-verify.",
            "Quality report chỉ chạy khi bật --enable-quality-report.",
            "Log production ghi JSON tối giản, output_rows có thể không được count để tránh scan.",
        ],
    )
    add_heading(doc, "9.2 Tránh small files", 2)
    add_bullets(
        doc,
        [
            "Survey snapshot là một Parquet tổng thay vì một file mỗi event.",
            "Kafka consumer gom tối đa 50 chat hoặc 15 giây rồi mới tạo JSONL Bronze.",
            "Silver/analytic dùng 4 output partitions; Gold summary dùng 1 file.",
            "Không partitionBy user_id hoặc timestamp có cardinality cao.",
            "Silver chỉ partition theo date/ingestion_date cần thiết.",
        ],
    )
    add_heading(doc, "9.3 Giảm shuffle và GCS I/O", 2)
    add_bullets(
        doc,
        [
            "AQE và adaptive coalesce bật.",
            "Không ép repartition analytic khi --analytic-compute-partitions=0; giữ natural Silver partitions.",
            "Bỏ temp compact stage mặc định.",
            "Gold core chạy riêng; question_distribution và numeric_summary là heavy mode, có thể chạy ít hơn.",
            "Gold survey dùng run_id mới thay vì overwrite toàn bộ root.",
        ],
    )
    add_heading(doc, "9.4 Tách core và heavy", 2)
    add_table(
        doc,
        ["Chế độ", "Bảng", "Tần suất đề xuất"],
        [
            ["Core", "analytic_features, overview, response_by_date, demographic", "Sau mỗi batch cần refresh Dashboard"],
            ["Heavy", "question_distribution, numeric_summary", "Hàng ngày hoặc khi cần phân tích sâu"],
        ],
        [1.2, 3.4, 1.8],
    )

    add_heading(doc, "10. Dashboard và logic Tổng quan", 1)
    add_bullets(
        doc,
        [
            "Dashboard chỉ đọc Gold.",
            "Survey Gold manifest tìm complete run mới nhất có đủ các bảng bắt buộc.",
            "Tổng quan dùng toàn bộ Học sinh + Sinh viên.",
            "Tab Học sinh và Sinh viên lọc theo Population/source_group.",
            "Chat Gold được lọc theo audience_group cho từng tab; Tổng quan đọc toàn bộ.",
            "KPI tỷ lệ nguy cơ survey dùng Target trung bình; chat không được cộng vào mẫu số survey.",
        ],
    )
    add_table(
        doc,
        ["Đối chiếu KPI Tổng quan", "Trước", "Sau"],
        [
            ["Tổng mẫu", "286.179", "286.280"],
            ["Số Target=1", "126.282", "126.333"],
            ["Tỷ lệ chính xác", "44,126928%", "44,129174%"],
            ["Hiển thị 2 số", "44,13%", "44,13%"],
            ["Hiển thị mới 3 số", "44,127%", "44,129%"],
        ],
        [2.5, 1.8, 1.8],
    )
    add_note(
        doc,
        "Giải thích",
        "101 survey mới chỉ chiếm tỷ trọng rất nhỏ trong tổng mẫu historical. Tỷ lệ chung tăng khoảng 0,002246 điểm %, nên khi làm tròn 2 số vẫn là 44,13%; hiển thị 3 số giúp nhìn thấy thay đổi.",
    )

    add_heading(doc, "11. Kết quả Dataproc thực tế", 1)
    add_table(
        doc,
        ["Batch ID", "Tầng", "Thời gian UTC", "Trạng thái"],
        [
            ["survey-up100-b2s-20260607091129", "Survey Bronze→Silver", "07/06/2026", "SUCCEEDED"],
            ["survey-up100-s2g-20260607091129", "Survey Silver→Gold core", "07/06/2026", "SUCCEEDED"],
            ["chat-up100-b2s-20260607091129", "Chat Bronze→Silver", "07/06/2026", "SUCCEEDED"],
            ["chat-up100-s2g-20260607091129", "Chat Silver→Gold", "07/06/2026", "SUCCEEDED"],
        ],
        [2.9, 1.7, 2.2, 1.0],
    )
    add_note(
        doc,
        "Đọc thời gian đúng",
        "Dataproc Serverless có thời gian khởi tạo runtime. Với dữ liệu hiện tại, phần lớn tổng thời gian batch gồm cả provisioning, không chỉ thời gian Spark transform/write.",
    )

    add_heading(doc, "12. Sự cố đã gặp và cách xử lý", 1)
    add_table(
        doc,
        ["Sự cố", "Nguyên nhân", "Xử lý", "Phòng ngừa"],
        [
            ["BE không lên", "Script mặc định PostgreSQL port 5433, PostgreSQL Ubuntu expose 5432", "Chạy BE với DATABASE_URL port 5432", "Chuẩn hóa env theo môi trường."],
            ["Kafka delivery timed out", "Kafka VM active nhưng BE local thiếu SSH tunnel 9092", "Khôi phục tunnel 127.0.0.1:9092 → VM localhost:9092", "Dùng service/tunnel monitor và health check trước test."],
            ["SSH mặc định bị từ chối", "Lệnh SSH không chỉ định khóa phù hợp", "Dùng ~/.ssh/google_compute_engine", "Chuẩn hóa SSH config và script tunnel."],
            ["Monitoring Alert Policy", "Ban đầu thiếu quyền; sau khi được cấp, payload uptime DELTA còn thiếu perSeriesAligner", "Thêm ALIGN_RATE/60s và tạo Kafka VM uptime alert thành công", "Giữ script tạo policy idempotent và kiểm tra notification channel."],
            ["KPI vẫn 44,13%", "Thay đổi nhỏ bị làm tròn 2 chữ số", "Đối chiếu Target và đổi KPI Tổng quan lên 3 chữ số", "Hiển thị thêm delta/run timestamp khi cần."],
            ["GCS có nhiều version Gold", "Versioned output giữ lịch sử run", "Dashboard chọn complete run mới nhất", "Thiết lập retention/lifecycle cho run cũ."],
        ],
        [1.6, 2.1, 2.0, 2.0],
    )

    add_heading(doc, "13. Runbook khởi chạy dịch vụ", 1)
    add_heading(doc, "13.1 PostgreSQL Ubuntu/WSL", 2)
    add_code(
        doc,
        """
sudo service postgresql start
sudo service postgresql status
""",
    )
    add_heading(doc, "13.2 Backend, Frontend và Dashboard", 2)
    add_code(
        doc,
        r"""
powershell -ExecutionPolicy Bypass -File scripts\deployment\run_backend.ps1
powershell -ExecutionPolicy Bypass -File scripts\deployment\run_frontend.ps1
powershell -ExecutionPolicy Bypass -File scripts\deployment\run_dashboard.ps1
""",
    )
    add_heading(doc, "13.3 Khởi tạo SSH tunnel Kafka cho BE local", 2)
    add_code(
        doc,
        r"""
ssh.exe -N `
  -o BatchMode=yes `
  -o ServerAliveInterval=30 `
  -o ExitOnForwardFailure=yes `
  -i $HOME\.ssh\google_compute_engine `
  -L 127.0.0.1:9092:localhost:9092 `
  Admin@34.21.211.62
""",
    )
    add_note(
        doc,
        "Bảo mật",
        "Không mở Kafka 9092 trực tiếp ra Internet. Tunnel chỉ là cách phục vụ Backend local. Khi Backend production nằm trong GCP/VPC, dùng private IP, firewall nội bộ và xác thực phù hợp.",
        "FFF3CD",
    )

    add_heading(doc, "14. Lệnh submit và chạy pipeline", 1)
    add_heading(doc, "14.1 Export survey pending vào Bronze snapshot", 2)
    add_code(
        doc,
        r"""
$env:DATABASE_URL="postgresql://student_app:<password-from-secret-manager>@127.0.0.1:5432/student_mental_health_app"
$env:SURVEY_KAFKA_ENABLED="false"
.\venv\Scripts\python.exe scripts\run_survey_snapshot_worker.py --once --limit 200
""",
    )
    add_heading(doc, "14.2 Chạy orchestration survey không Kafka", 2)
    add_code(
        doc,
        r"""
powershell -ExecutionPolicy Bypass -File scripts\scheduler\survey_no_kafka_dashboard_refresh.ps1 `
  -SkipSurveySubmit
""",
    )
    add_heading(doc, "14.3 Chạy nightly đầy đủ survey + chat", 2)
    add_code(
        doc,
        r"""
powershell -ExecutionPolicy Bypass -File scripts\scheduler\nightly_dashboard_refresh.ps1 `
  -ProcessDate "2026-06-06"

# Khi cần chạy thêm bảng survey nặng:
powershell -ExecutionPolicy Bypass -File scripts\scheduler\nightly_dashboard_refresh.ps1 `
  -ProcessDate "2026-06-06" `
  -IncludeHeavySurveyTables
""",
    )
    add_heading(doc, "14.4 Ví dụ gcloud submit trực tiếp", 2)
    add_code(
        doc,
        """
gcloud dataproc batches submit pyspark \
  gs://student-mental-health-lake-nhom1-2026/scripts/survey_pipeline/survey_bronze_to_silver_spark.py \
  --project=student-mental-health-496205 \
  --region=asia-southeast1 \
  --batch=survey-b2s-<timestamp> \
  --properties=spark.dynamicAllocation.enabled=false,spark.executor.instances=2,spark.executor.cores=4,spark.executor.memory=8g,spark.driver.cores=4,spark.driver.memory=8g,spark.default.parallelism=8,spark.sql.shuffle.partitions=8,spark.sql.adaptive.enabled=true,spark.sql.adaptive.coalescePartitions.enabled=true,spark.hadoop.mapreduce.fileoutputcommitter.algorithm.version=2 \
  -- --fast-mode --output-partitions 4 --spark-parallelism 8 --shuffle-partitions 8
""",
    )
    add_heading(doc, "14.5 Chạy kiểm thử 100 tài khoản", 2)
    add_code(
        doc,
        r"""
.\venv\Scripts\python.exe scripts\e2e\submit_test_surveys_batch.py `
  --api-base-url http://127.0.0.1:8000 `
  --count 100 `
  --workers 5 `
  --chat-messages-per-user 1 `
  --email-prefix survey.chat.upgrade100.07062026 `
  --output-json logs\survey_chat_upgrade_e2e_100_20260607160356.json
""",
    )

    add_heading(doc, "15. Lệnh kiểm tra trạng thái và xem log", 1)
    add_heading(doc, "15.1 Kiểm tra dịch vụ local", 2)
    add_code(
        doc,
        r"""
Invoke-WebRequest -UseBasicParsing http://127.0.0.1:8000/health
Invoke-WebRequest -UseBasicParsing http://127.0.0.1:5173
Invoke-WebRequest -UseBasicParsing http://127.0.0.1:8501
Test-NetConnection 127.0.0.1 -Port 9092
Get-NetTCPConnection -State Listen | Where-Object {$_.LocalPort -in 8000,5173,8501,9092}
""",
    )
    add_heading(doc, "15.2 Kiểm tra Kafka VM", 2)
    add_code(
        doc,
        r"""
ssh -i $HOME\.ssh\google_compute_engine Admin@34.21.211.62 `
  "systemctl is-active kafka student-chat-kafka-consumer"

ssh -i $HOME\.ssh\google_compute_engine Admin@34.21.211.62 `
  "sudo journalctl -u kafka -u student-chat-kafka-consumer --since '30 minutes ago' --no-pager"

ssh -i $HOME\.ssh\google_compute_engine Admin@34.21.211.62 `
  "sudo /opt/kafka/bin/kafka-topics.sh --bootstrap-server localhost:9092 --describe --topic student-chat-logs"
""",
    )
    add_heading(doc, "15.3 Kiểm tra Dataproc", 2)
    add_code(
        doc,
        """
gcloud dataproc batches list --project=student-mental-health-496205 --region=asia-southeast1
gcloud dataproc batches describe <batch-id> --project=student-mental-health-496205 --region=asia-southeast1
gcloud dataproc batches wait <batch-id> --project=student-mental-health-496205 --region=asia-southeast1
""",
    )
    add_heading(doc, "15.4 Kiểm tra GCS", 2)
    add_code(
        doc,
        """
gcloud storage ls --recursive gs://student-mental-health-lake-nhom1-2026/bronze/app_survey_snapshot/
gcloud storage ls --recursive gs://student-mental-health-lake-nhom1-2026/bronze/chat_logs/date=2026-06-06/
gcloud storage ls --recursive gs://student-mental-health-lake-nhom1-2026/silver/survey_cleaned/
gcloud storage ls --recursive gs://student-mental-health-lake-nhom1-2026/gold/dashboard_tables/
""",
    )
    add_heading(doc, "15.5 Đọc log local", 2)
    add_code(
        doc,
        r"""
Get-ChildItem logs -File | Sort-Object LastWriteTime -Descending | Select-Object -First 30
Get-Content logs\dataproc_survey_e2e50_20260606153018.log
Get-Content logs\dataproc_chat_e2e50_20260606154711.log
Get-Content logs\backend_50e2e_direct_20260606222303.err.log -Tail 100
""",
    )

    add_heading(doc, "16. Cron và lịch chạy đề xuất", 1)
    add_table(
        doc,
        ["Công việc", "Lịch đề xuất", "Ghi chú"],
        [
            ["Survey snapshot worker", "Mỗi 5-15 phút hoặc event-driven", "Single writer; limit theo backlog."],
            ["Survey Gold core", "Mỗi 30-60 phút hoặc sau snapshot batch", "Phục vụ Dashboard nhanh."],
            ["Survey Gold heavy", "01:30 hàng ngày", "Question distribution/numeric summary."],
            ["Chat Bronze→Silver backfill", "02:00 hàng ngày", "Streaming/consumer vẫn chạy liên tục; batch là lớp backfill."],
            ["Chat Silver→Gold", "02:15 hàng ngày hoặc mỗi giờ", "Tùy yêu cầu độ mới Dashboard."],
            ["GCS lifecycle cleanup", "Hàng ngày theo policy", "Xóa run Gold cũ sau thời gian giữ lại."],
        ],
        [2.1, 1.9, 2.5],
    )
    add_note(
        doc,
        "Khung giờ",
        "Chọn giờ ít người hoạt động theo múi giờ Asia/Saigon. Dataproc Serverless không cần giữ cluster luôn chạy; chỉ phát sinh tài nguyên khi submit batch.",
    )

    add_heading(doc, "17. Quyền IAM cần thiết", 1)
    add_table(
        doc,
        ["Đối tượng", "Quyền/role tối thiểu tham khảo", "Mục đích"],
        [
            ["Người submit Dataproc", "roles/dataproc.editor; roles/iam.serviceAccountUser", "Tạo, xem và chạy Dataproc Serverless batch."],
            ["Dataproc service account", "roles/dataproc.worker; quyền đọc/ghi đúng GCS prefix", "Đọc scripts/Bronze/Silver và ghi Silver/Gold."],
            ["Snapshot worker/consumer", "roles/storage.objectAdmin trên bucket hoặc prefix tương ứng", "Đọc/ghi snapshot survey và chat Bronze."],
            ["Dashboard", "roles/storage.objectViewer", "Chỉ đọc Gold."],
            ["Quản trị Kafka VM", "roles/compute.instanceAdmin.v1; roles/compute.osAdminLogin nếu dùng OS Login", "Quản lý VM và SSH."],
            ["Xem log", "roles/logging.viewer", "Đọc Cloud Logging."],
            ["Quản trị cảnh báo", "roles/monitoring.alertPolicyEditor", "Tạo/sửa Cloud Monitoring Alert Policy. Quyền đã có hiệu lực và Kafka VM uptime alert đã được tạo."],
        ],
        [1.8, 3.1, 1.8],
    )
    add_note(
        doc,
        "Least privilege",
        "Production nên cấp quyền theo service account và prefix/bucket cần thiết, không dùng Owner/Editor toàn project nếu không bắt buộc.",
        "FFF3CD",
    )

    add_heading(doc, "18. Checklist vận hành production", 1)
    add_bullets(
        doc,
        [
            "Kiểm tra PostgreSQL active và Backend health trả status ok.",
            "Kiểm tra Kafka broker + consumer active; nếu BE local thì kiểm tra tunnel 9092.",
            "Kiểm tra survey snapshot unique user_id và không có schema lỗi.",
            "Submit Bronze→Silver trước; chỉ chạy Silver→Gold sau khi batch trước SUCCEEDED.",
            "Kiểm tra JSON log cuối job: status, tables_written, skipped_tables, partition_config và duration.",
            "Đối chiếu Gold overview với analytic_features và số lượng theo source_group.",
            "Restart/clear cache Dashboard khi cần kiểm tra ngay sau batch.",
            "Không bật quality report/counts trong production thường xuyên.",
            "Không partitionBy user_id/timestamp và không tăng partition tùy tiện.",
            "Không ghi báo cáo hoặc file không phải dữ liệu xử lý vào GCS Data Lake.",
        ],
    )

    add_heading(doc, "19. Nâng cấp đã thực hiện và khuyến nghị tiếp theo", 1)
    add_table(
        doc,
        ["Nâng cấp", "Trạng thái", "Kết quả"],
        [
            ["Kafka producer idempotent + acks=all + retry", "Đã hoàn thành", "Giảm nguy cơ mất/trùng event khi mạng chập chờn."],
            ["Kafka tunnel watchdog", "Đã hoàn thành", "scripts/deployment/ensure_kafka_tunnel.ps1 kiểm tra và mở tunnel khi cần."],
            ["Health check production", "Đã hoàn thành", "Kiểm tra BE, FE, Dashboard, tunnel, Survey/Chat Bronze và Gold."],
            ["Survey Gold current manifest", "Đã hoàn thành", "Dashboard đọc trực tiếp run Gold hiện hành survey_upgrade100_20260607091129."],
            ["Chuẩn hóa PostgreSQL port", "Đã hoàn thành", "Script khởi động BE mặc định dùng PostgreSQL Ubuntu cổng 5432."],
            ["Cloud Monitoring Kafka VM uptime alert", "Đã hoàn thành", "Policy 17054748573461416323 cảnh báo khi VM không báo uptime trong 5 phút."],
            ["Kafka HA nhiều broker / managed Kafka", "Chưa triển khai", "Cần Backend trong VPC hoặc kết nối nhiều endpoint trước khi triển khai HA thực sự."],
        ],
        [3.0, 1.4, 3.0],
    )
    add_table(
        doc,
        ["Mức ưu tiên", "Khuyến nghị", "Lợi ích"],
        [
            ["Cao", "Chuyển Kafka single-node sang managed Kafka hoặc tối thiểu 3 broker", "HA và replication thực sự."],
            ["Cao", "Chạy Backend production trong GCP/VPC thay cho tunnel local", "Kết nối Kafka ổn định, bảo mật hơn."],
            ["Trung bình", "GCS lifecycle cho Gold versioned runs và temp", "Giảm chi phí và object clutter."],
            ["Trung bình", "Mở rộng Cloud Monitoring từ VM uptime sang Kafka lag, Dataproc failure và snapshot age", "Phát hiện lỗi sớm hơn ở từng tầng."],
            ["Thấp", "Benchmark tự động core/heavy và partition 4/8 theo kích thước dữ liệu", "Tự điều chỉnh khi dữ liệu tăng lớn."],
        ],
        [1.0, 3.7, 2.2],
    )

    add_heading(doc, "20. Kết luận", 1)
    doc.add_paragraph(
        "Luồng Survey và Chat hiện đã hoạt động end-to-end từ Web/Backend đến GCS Bronze, Spark Silver/Gold và Dashboard. "
        "Các tối ưu quan trọng đã áp dụng đúng với quy mô dữ liệu hiện tại: giảm Spark actions, partition vừa đủ, AQE, Parquet, "
        "gom batch chống small files, tách core/heavy, versioned Gold và đọc complete run mới nhất. Bốn Dataproc batch kiểm thử "
        "đều SUCCEEDED, snapshot survey không duplicate, Kafka consumer ghi GCS thành công và Dashboard phản ánh dữ liệu mới."
    )
    doc.add_paragraph(
        "Điểm cần ưu tiên trước production quy mô lớn là nâng Kafka từ single-node lên kiến trúc HA, đưa Backend vào VPC phù hợp và "
        "mở rộng monitoring từ VM uptime sang Kafka lag, Dataproc failure và snapshot age. Producer idempotent, tunnel watchdog, production "
        "health check, Kafka VM uptime alert và Survey Gold current manifest đã được triển khai mà không thay đổi business logic hiện tại."
    )

    add_heading(doc, "Phụ lục A. File code chính", 1)
    add_table(
        doc,
        ["File", "Vai trò"],
        [
            ["MentalSchool_Dashboard/app.py", "Dashboard đọc Gold và hiển thị Tổng quan/Học sinh/Sinh viên."],
            ["MentalSchool_Dashboard/dashboard_gcs_loader.py", "Loader và chuẩn hóa bảng Gold chat."],
            ["MentalSchool_Dashboard/survey_bronze_to_silver_spark.py", "Survey Bronze→Silver."],
            ["MentalSchool_Dashboard/survey_silver_to_gold_spark.py", "Survey Silver→Gold."],
            ["MentalSchool_Dashboard/chat_bronze_to_silver_spark.py", "Chat Bronze→Silver."],
            ["MentalSchool_Dashboard/chat_silver_to_gold_spark.py", "Chat Silver→Gold."],
            ["backend/api/survey.py", "Survey API."],
            ["backend/surveys/snapshot.py", "Merge/validate/temp-write snapshot Parquet."],
            ["backend/surveys/snapshot_worker.py", "Single writer và PostgreSQL advisory lock."],
            ["backend/chat_logs/kafka_publisher.py", "Chat producer Kafka với metadata profile."],
            ["scripts/kafka/run_kafka_consumer.py", "Kafka consumer ghi GCS Bronze."],
            ["scripts/kafka/setup_chat_kafka_vm.sh", "Cấu hình Kafka VM và systemd consumer."],
            ["scripts/scheduler/nightly_dashboard_refresh.ps1", "Orchestration nightly survey + chat."],
            ["scripts/scheduler/survey_no_kafka_dashboard_refresh.ps1", "Survey snapshot + Dataproc không Kafka."],
            ["scripts/deployment/ensure_kafka_tunnel.ps1", "Kiểm tra và khôi phục SSH Kafka tunnel."],
            ["scripts/deployment/production_health_check.py", "Health check toàn bộ luồng production."],
            ["scripts/deployment/create_kafka_monitoring_alert.py", "Tạo idempotent Kafka VM uptime alert khi quyền Monitoring có hiệu lực."],
            ["scripts/update_survey_gold_current_manifest.py", "Cập nhật pointer Survey Gold hiện hành cho Dashboard."],
            ["scripts/e2e/submit_test_surveys_batch.py", "Kiểm thử nhiều tài khoản survey/chat."],
            ["frontend/src/main.jsx", "Survey UX, draft, khóa/mở Dashboard, nhãn đáp án."],
        ],
        [3.8, 3.0],
    )

    add_heading(doc, "Phụ lục B. Log kiểm thử quan trọng", 1)
    add_bullets(
        doc,
        [
            "logs/survey_chat_upgrade_e2e_100_20260607160356.json",
            "logs/survey_chat_upgrade_e2e_100_20260607160356.run.log",
            "logs/dataproc_upgrade100_20260607091129.log",
            "logs/survey_gold_manifest_upgrade100.log",
            "logs/backend_upgrade100_20260607160148.err.log",
        ],
    )

    add_heading(doc, "Phụ lục C. Danh mục file code và mục đích sử dụng", 1)
    add_note(
        doc,
        "Cách đọc",
        "Giữ = đang nằm trong luồng chính. Giữ khi cần = setup, backfill, kiểm thử hoặc khôi phục. Ứng viên dọn dẹp = không nằm trong scheduler/luồng hiện tại; chỉ xóa sau khi xác minh không còn người dùng hoặc tài liệu phụ thuộc.",
    )
    add_heading(doc, "C.1 File production cần giữ", 2)
    add_table(
        doc,
        ["File/nhóm file", "Dùng để làm gì", "Kết luận"],
        [
            ["backend/main.py", "Khởi tạo FastAPI, gắn router và schema database.", "Giữ"],
            ["backend/api/rag.py", "API chatbot; lấy profile và publish chat event sang Kafka.", "Giữ"],
            ["backend/api/survey.py", "API trạng thái, câu hỏi, hoãn và submit survey.", "Giữ"],
            ["backend/api/users.py + schemas.py", "Đăng ký, đăng nhập, profile và schema API.", "Giữ"],
            ["backend/db/connection.py + schema.sql", "Kết nối PostgreSQL và tạo schema.", "Giữ"],
            ["backend/db/users.py + surveys.py + chat_sessions.py", "Lưu user/profile/survey và ánh xạ phiên chat.", "Giữ"],
            ["backend/surveys/questions.py + definitions.py + mapping.py", "Đọc định nghĩa câu hỏi, validate và map đáp án sang cột dashboard.", "Giữ"],
            ["backend/surveys/snapshot.py + snapshot_worker.py", "Tạo snapshot Parquet an toàn, unique user_id và single writer.", "Giữ"],
            ["backend/chat_logs/kafka_publisher.py + gcs_writer.py", "Tạo chat event, ẩn danh/mask PII và gửi Kafka.", "Giữ"],
            ["backend/rag/**", "Pipeline RAG, embeddings, retrieval, emotion và safety cho chatbot.", "Giữ nếu chatbot còn dùng"],
            ["frontend/src/main.jsx + App.css", "Toàn bộ giao diện Web, survey wizard, chatbot và liên kết Dashboard.", "Giữ"],
            ["frontend/package.json + package-lock.json + index.html", "Cấu hình build và dependency Frontend.", "Giữ"],
            ["MentalSchool_Dashboard/app.py", "Ứng dụng Streamlit Dashboard.", "Giữ"],
            ["MentalSchool_Dashboard/dashboard_gcs_loader.py", "Đọc và chuẩn hóa Gold chat cho Dashboard.", "Giữ"],
            ["MentalSchool_Dashboard/assets/styles.css", "CSS của Dashboard.", "Giữ"],
            ["MentalSchool_Dashboard/survey_bronze_to_silver_spark.py", "Spark Survey Bronze→Silver.", "Giữ"],
            ["MentalSchool_Dashboard/survey_silver_to_gold_spark.py", "Spark Survey Silver→Gold.", "Giữ"],
            ["MentalSchool_Dashboard/chat_bronze_to_silver_spark.py", "Spark Chat Bronze→Silver/backfill.", "Giữ"],
            ["MentalSchool_Dashboard/chat_silver_to_gold_spark.py", "Spark Chat Silver→Gold.", "Giữ"],
            ["MentalSchool_Dashboard/requirements.txt", "Dependency chạy Dashboard/Spark code.", "Giữ"],
            ["school_survey_questions_answers.txt", "Nguồn định nghĩa khảo sát Học sinh.", "Giữ"],
            ["university_survey_questions_answers.txt", "Nguồn định nghĩa khảo sát Sinh viên.", "Giữ"],
        ],
        [3.4, 3.2, 1.2],
    )

    add_heading(doc, "C.2 File vận hành, setup và backfill", 2)
    add_table(
        doc,
        ["File", "Dùng để làm gì", "Kết luận"],
        [
            ["scripts/deployment/run_backend.ps1/.bat", "Khởi chạy Backend local.", "Giữ khi cần"],
            ["scripts/deployment/run_frontend.ps1/.bat", "Khởi chạy Frontend local.", "Giữ khi cần"],
            ["scripts/deployment/run_dashboard.ps1/.bat", "Khởi chạy Dashboard local.", "Giữ khi cần"],
            ["scripts/deployment/run_all.bat", "Mở PostgreSQL, BE, FE và Dashboard cùng lúc.", "Giữ khi cần"],
            ["scripts/deployment/run_postgres.bat", "Khởi chạy PostgreSQL theo cấu hình local.", "Kiểm tra port trước khi dùng"],
            ["scripts/deployment/gcs_login.ps1/.bat", "Đăng nhập GCP và kiểm tra quyền GCS.", "Giữ khi cần"],
            ["scripts/deployment/ensure_kafka_tunnel.ps1", "Kiểm tra cổng 9092 và mở lại SSH Kafka tunnel khi thiếu.", "Giữ"],
            ["scripts/deployment/production_health_check.py", "Kiểm tra health BE/FE/Dashboard/Kafka tunnel và độ sẵn sàng Bronze/Gold.", "Giữ"],
            ["scripts/deployment/create_kafka_monitoring_alert.py", "Tạo Kafka VM uptime alert; yêu cầu roles/monitoring.alertPolicyEditor.", "Giữ"],
            ["scripts/deployment/upload_dashboard_pipeline_scripts.ps1", "Upload Spark/Kafka/scheduler scripts lên GCS.", "Giữ"],
            ["scripts/scheduler/nightly_dashboard_refresh.ps1", "Submit Survey + Chat Dataproc theo lịch.", "Giữ"],
            ["scripts/scheduler/survey_no_kafka_dashboard_refresh.ps1", "Export snapshot và refresh Survey khi chưa dùng Kafka survey.", "Giữ"],
            ["scripts/run_survey_snapshot_worker.py", "Chạy single snapshot worker một lần hoặc theo lịch.", "Giữ"],
            ["scripts/rebuild_app_survey_snapshot.py", "Rebuild toàn bộ survey snapshot từ PostgreSQL khi khôi phục/backfill.", "Giữ khi cần"],
            ["scripts/export_postgres_snapshots_to_gcs.py", "Export snapshot PostgreSQL sang GCS cho backfill/audit.", "Giữ khi cần"],
            ["scripts/standardize_legacy_surveys.py", "Chuẩn hóa lại dữ liệu historical survey.", "Giữ khi cần"],
            ["scripts/build_dashboard_gold_from_standardized.py", "Tạo/kiểm tra Gold từ standardized trong giai đoạn chuyển đổi.", "Giữ khi cần"],
            ["scripts/update_survey_gold_current_manifest.py", "Tạo current manifest trỏ đến Survey Gold complete run mới nhất.", "Giữ"],
            ["scripts/kafka/setup_chat_kafka_vm.sh", "Cài Kafka KRaft và consumer systemd trên VM.", "Giữ"],
            ["scripts/kafka/run_kafka_consumer.py", "Consumer chat Kafka → GCS Bronze.", "Giữ"],
            ["scripts/kafka/run_survey_snapshot_consumer.py", "Consumer survey Kafka trong tương lai.", "Giữ khi tích hợp Kafka survey"],
            ["scripts/preprocessing/**", "Tiền xử lý knowledge base/RAG trên Dataproc.", "Giữ nếu RAG còn dùng"],
            ["scripts/embeddings/**", "Export/index/test Qdrant embeddings.", "Giữ nếu RAG/Qdrant còn dùng"],
            ["scripts/run_rag_pipeline.py", "Điều phối pipeline RAG.", "Giữ nếu RAG còn dùng"],
        ],
        [3.6, 3.2, 1.3],
    )

    add_heading(doc, "C.3 File kiểm thử và tài liệu", 2)
    add_table(
        doc,
        ["File/nhóm file", "Dùng để làm gì", "Kết luận"],
        [
            ["scripts/e2e/submit_test_survey.py", "Test một tài khoản submit survey.", "Giữ cho kiểm thử"],
            ["scripts/e2e/submit_test_surveys_batch.py", "Test nhiều tài khoản survey + chat.", "Giữ cho kiểm thử tải"],
            ["scripts/deployment/seed_demo_users.py", "Tạo user demo.", "Xóa nếu production không cần demo"],
            ["scripts/generate_big_data_report_07062026.py", "Sinh lại báo cáo Word này.", "Giữ nếu cần cập nhật báo cáo"],
            ["test_qdrant_connection.py", "Kiểm tra kết nối Qdrant.", "Giữ nếu RAG còn dùng"],
            ["MentalSchool_Dashboard/test_gcs_connection.py", "Kiểm tra đọc GCS.", "Giữ khi debug"],
            ["MentalSchool_Dashboard/test_gcs_write.py", "Kiểm tra quyền ghi GCS.", "Xóa/khóa ở production nếu không cần"],
            ["MentalSchool_Dashboard/read_bronze_sample.py", "Đọc mẫu Bronze để debug.", "Giữ khi debug"],
            ["MentalSchool_Dashboard/dashboard_columns_questions.txt", "Tham chiếu cột/câu hỏi Dashboard.", "Giữ làm tài liệu"],
            ["*.md báo cáo/runbook và Bao_cao_big_data.07062026.docx", "Tài liệu vận hành và lịch sử quyết định.", "Giữ local; không upload GCS Data Lake"],
            ["logs/** và ảnh kiểm thử", "Bằng chứng chạy/test cục bộ.", "Có thể archive/xóa theo retention"],
            ["frontend/dist/**", "Build output Frontend.", "Có thể xóa và npm run build lại; giữ nếu deploy static trực tiếp"],
            ["**/__pycache__/**", "Python bytecode cache.", "Có thể xóa an toàn"],
        ],
        [3.7, 3.0, 1.5],
    )

    add_heading(doc, "C.4 Ứng viên dọn dẹp sau khi xác minh", 2)
    add_table(
        doc,
        ["File", "Lý do cần xem xét", "Hành động đề xuất"],
        [
            ["MentalSchool_Dashboard/bronze_to_silver.py", "Tên pipeline cũ, không được scheduler hiện tại gọi.", "So sánh rồi archive/xóa nếu không còn dùng"],
            ["MentalSchool_Dashboard/bronze_to_silver_safe.py", "Biến thể cũ, không được scheduler hiện tại gọi.", "So sánh rồi archive/xóa"],
            ["MentalSchool_Dashboard/bronze_to_silver_spark.py", "Pipeline Spark chung cũ; pipeline hiện tại dùng survey/chat riêng.", "Xác minh trước khi xóa"],
            ["MentalSchool_Dashboard/silver_to_gold.py", "Tên pipeline cũ, scheduler hiện tại không gọi.", "So sánh rồi archive/xóa"],
            ["MentalSchool_Dashboard/silver_to_gold_safe.py", "Biến thể cũ, scheduler hiện tại không gọi.", "So sánh rồi archive/xóa"],
            ["MentalSchool_Dashboard/chat_streaming_to_gold_spark.py", "Không nằm trong nightly scheduler hiện tại.", "Giữ nếu còn streaming Gold; nếu không thì archive"],
            ["MentalSchool_Dashboard/chat_kafka_to_silver_streaming.py", "Không nằm trong batch nightly; có thể dùng khi triển khai Spark Structured Streaming.", "Chỉ xóa nếu xác nhận consumer JSONL là luồng duy nhất"],
            ["MentalSchool_Dashboard/big-data-processing_updated.ipynb", "Notebook nghiên cứu, không phải production runtime.", "Archive ngoài source production"],
            ["MentalSchool_Dashboard/gcs_client.py", "Cần kiểm tra import; app hiện dùng loader/client nội bộ khác.", "Xác minh rồi xóa nếu không import"],
            ["MentalSchool_Dashboard/utils/dashboard_core.py", "Cần kiểm tra import và chức năng cũ.", "Xác minh rồi xóa nếu không import"],
            ["MentalSchool_Dashboard/run_dashboard.bat", "Trùng vai trò với scripts/deployment/run_dashboard.*.", "Giữ một cách chạy chuẩn"],
            ["test-permission.txt", "File thử quyền, không phải code/runtime.", "Có thể xóa"],
        ],
        [3.8, 3.0, 1.8],
    )
    add_note(
        doc,
        "Trước khi xóa",
        "Tạo branch/commit hoặc archive; dùng rg để tìm import/tham chiếu; kiểm tra scheduler, GCS scripts và tài liệu; chạy lại BE/FE/Dashboard cùng pipeline E2E. Không xóa file chỉ vì tên cũ nếu chưa đối chiếu nội dung.",
        "FFF3CD",
    )
    return doc


def main() -> None:
    doc = build_document()
    try:
        doc.save(OUTPUT)
        print(OUTPUT)
    except PermissionError:
        fallback = OUTPUT.with_name(f"{OUTPUT.stem}.updated{OUTPUT.suffix}")
        doc.save(fallback)
        print(f"{OUTPUT} đang được mở; đã lưu báo cáo cập nhật tại {fallback}")


if __name__ == "__main__":
    main()
